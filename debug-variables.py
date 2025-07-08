#!/usr/bin/env python3

"""
Script para debugar variáveis no template DOCX
Execute: python debug-variables.py
"""

from docx import Document
import re

def debug_template_variables():
    template_path = "CONTRATO Casa da Ana.docx"
    
    try:
        document = Document(template_path)
        print("🔍 DEBUGANDO VARIÁVEIS NO TEMPLATE")
        print("=" * 50)
        
        # Lista para armazenar todas as variáveis encontradas
        variables_found = set()
        
        print("\n📄 ANALISANDO PARÁGRAFOS:")
        print("-" * 30)
        
        for i, paragraph in enumerate(document.paragraphs):
            text = paragraph.text.strip()
            if text and '{{' in text:
                print(f"Parágrafo {i}: {text}")
                
                # Encontrar todas as variáveis {{ }}
                variables = re.findall(r'\{\{[^}]+\}\}', text)
                for var in variables:
                    variables_found.add(var)
                    print(f"  → Variável encontrada: '{var}'")
        
        print("\n📊 ANALISANDO TABELAS:")
        print("-" * 30)
        
        for table_idx, table in enumerate(document.tables):
            print(f"\nTabela {table_idx}:")
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        text = paragraph.text.strip()
                        if text and '{{' in text:
                            print(f"  Linha {row_idx}, Célula {cell_idx}: {text}")
                            
                            variables = re.findall(r'\{\{[^}]+\}\}', text)
                            for var in variables:
                                variables_found.add(var)
                                print(f"    → Variável encontrada: '{var}'")
        
        print("\n✅ RESUMO DAS VARIÁVEIS ENCONTRADAS:")
        print("=" * 50)
        
        if variables_found:
            for var in sorted(variables_found):
                print(f"  • {var}")
        else:
            print("  ❌ Nenhuma variável {{ }} encontrada!")
        
        print(f"\n📊 Total de variáveis únicas: {len(variables_found)}")
        
        # Verificar variáveis específicas que deveriam estar presentes
        expected_vars = [
            '{{ nome do locatário }}',
            '{{ qtd_noites }}', 
            '{{ dia_inicio }}',
            '{{ dia_fim }}',
            '{{ valor_locacao }}',
            '{{ valor_locacao/2 }}'
        ]
        
        print("\n🔍 VERIFICAÇÃO DE VARIÁVEIS ESPERADAS:")
        print("-" * 40)
        
        for expected in expected_vars:
            if expected in variables_found:
                print(f"  ✅ {expected}")
            else:
                # Procurar variações próximas
                similar = [v for v in variables_found if expected.replace('{{ ', '').replace(' }}', '') in v.replace('{{ ', '').replace(' }}', '')]
                if similar:
                    print(f"  ⚠️  {expected} → Encontrada variação: {similar[0]}")
                else:
                    print(f"  ❌ {expected} → NÃO ENCONTRADA")
        
        print("\n💡 DICAS PARA CORREÇÃO:")
        print("-" * 25)
        print("  1. Verifique se as variáveis no DOCX têm exatamente a mesma grafia")
        print("  2. Cuidado com espaços extras dentro das chaves")
        print("  3. Verifique se não há quebras de linha dentro das variáveis")
        print("  4. Use as variáveis exatas encontradas acima no código Python")
        
    except FileNotFoundError:
        print(f"❌ Arquivo '{template_path}' não encontrado!")
        print("  Verifique se o template está no diretório atual")
    except Exception as e:
        print(f"❌ Erro ao processar template: {e}")

if __name__ == "__main__":
    debug_template_variables() 