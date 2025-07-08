import os
import base64
import requests
from docx import Document
from datetime import datetime
import logging

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ContractService:
    def __init__(self, webhook_url=None):
        """
        Inicializa o serviço de contrato com configurações do webhook
        """
        self.webhook_url = webhook_url or os.getenv('WEBHOOK_URL', 'https://webh.criativamaisdigital.com.br/webhook/c1d01bf8-6d34-44ee-9100-2923b5fb7876')
        self.template_path = "CONTRATO Casa da Ana.docx"
    
    def validate_contract_data(self, dados):
        """
        Valida se todos os dados obrigatórios foram fornecidos
        """
        required_fields = [
            'nome_do_locatario',
            'estado_civil',
            'nacionalidade',
            'profissao',
            'numero_do_rg',
            'numero_do_cpf',
            'telefone_celular',
            'email',
            'endereco',
            'qtd_noites',
            'dia_inicio',
            'dia_fim',
            'valor_locacao'
        ]
        
        missing_fields = []
        for field in required_fields:
            if not dados.get(field):
                missing_fields.append(field)
        
        if missing_fields:
            raise ValueError(f"Campos obrigatórios não fornecidos: {', '.join(missing_fields)}")
        
        return True
    
    def generate_contract(self, dados_locatario):
        """
        Gera o contrato preenchendo o template com os dados fornecidos
        """
        try:
            self.validate_contract_data(dados_locatario)
            
            # Verifica se o template existe
            if not os.path.exists(self.template_path):
                raise FileNotFoundError(f"Template '{self.template_path}' não encontrado")
            
            # Converte os dados para o formato esperado pelo template
            dados_template = {
                # Variáveis originais
                '{{ nome do locatário }}': dados_locatario['nome_do_locatario'],
                '{{ estado civil }}': dados_locatario['estado_civil'],
                '{{ nacionalidade }}': dados_locatario['nacionalidade'],
                '{{ profissao }}': dados_locatario['profissao'],
                '{{ numero do rg }}': dados_locatario['numero_do_rg'],
                '{{ numero do cpf }}': dados_locatario['numero_do_cpf'],
                '{{ telefone celular }}': dados_locatario['telefone_celular'],
                '{{ e-mail }}': dados_locatario['email'],
                '{{ endereco }}': dados_locatario['endereco'],
                
                # Novas variáveis - testando diferentes formatos
                '{{ qtd_noites }}': str(dados_locatario['qtd_noites']),
                '{{qtd_noites}}': str(dados_locatario['qtd_noites']),  # Sem espaços
                '{{ qtd noites }}': str(dados_locatario['qtd_noites']),  # Com espaço
                
                '{{ dia_inicio }}': dados_locatario['dia_inicio'],
                '{{dia_inicio}}': dados_locatario['dia_inicio'],  # Sem espaços
                '{{ dia inicio }}': dados_locatario['dia_inicio'],  # Com espaço
                
                '{{ dia_fim }}': dados_locatario['dia_fim'],
                '{{dia_fim}}': dados_locatario['dia_fim'],  # Sem espaços
                '{{ dia fim }}': dados_locatario['dia_fim'],  # Com espaço
                
                '{{ valor_locacao }}': dados_locatario['valor_locacao'],
                '{{valor_locacao}}': dados_locatario['valor_locacao'],  # Sem espaços
                '{{ valor locacao }}': dados_locatario['valor_locacao'],  # Com espaço
                
                '{{ valor_locacao/2 }}': self._calculate_half_value(dados_locatario['valor_locacao']),
                '{{valor_locacao/2}}': self._calculate_half_value(dados_locatario['valor_locacao']),  # Sem espaços
                '{{ valor locacao/2 }}': self._calculate_half_value(dados_locatario['valor_locacao']),  # Com espaço
            }
            
            # Gera nome do arquivo
            nome_sanitizado = self._sanitize_filename(dados_locatario['nome_do_locatario'])
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"CONTRATO_{nome_sanitizado}_{timestamp}.docx"
            
            # Preenche o contrato
            success = self._fill_contract_template(dados_template, output_filename)
            
            if success:
                return output_filename
            else:
                raise Exception("Falha ao preencher o template do contrato")
                
        except Exception as e:
            logger.error(f"Erro ao gerar contrato: {str(e)}")
            raise
    
    def _sanitize_filename(self, nome):
        """
        Sanitiza o nome para uso em nome de arquivo
        """
        nome_limpo = "".join(c if c.isalnum() or c in (' ', '_') else '_' for c in nome).rstrip()
        return nome_limpo.replace(' ', '_')
    
    def _calculate_half_value(self, valor_string):
        """
        Calcula metade do valor da locação, tratando diferentes formatos
        """
        try:
            # Remove símbolos e converte para float
            valor_limpo = valor_string
            if isinstance(valor_limpo, str):
                # Remove R$, pontos de milhares e converte vírgula para ponto
                valor_limpo = valor_limpo.replace('R$', '').replace(' ', '').strip()
                # Se tem vírgula, trata como separador decimal brasileiro
                if ',' in valor_limpo:
                    if valor_limpo.count('.') > 0 and valor_limpo.count(',') == 1:
                        # Formato: 1.000,50 (brasileiro)
                        valor_limpo = valor_limpo.replace('.', '').replace(',', '.')
                    else:
                        # Formato: 1000,50
                        valor_limpo = valor_limpo.replace(',', '.')
                
                valor_float = float(valor_limpo)
            else:
                valor_float = float(valor_string)
            
            metade = valor_float / 2
            
            # Formata de volta para o padrão brasileiro
            return f"R$ {metade:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            
        except (ValueError, AttributeError) as e:
            logger.warning(f"Erro ao calcular metade do valor '{valor_string}': {e}")
            return "R$ 0,00"
    
    def _fill_contract_template(self, dados_template, output_filename):
        """
        Preenche o template DOCX com os dados fornecidos
        """
        try:
            document = Document(self.template_path)
            substituicoes_realizadas = {}

            # Itera sobre todos os parágrafos do documento
            for paragraph in document.paragraphs:
                original_text = paragraph.text
                if '{{' in original_text:
                    logger.info(f"Parágrafo com variáveis encontrado: {original_text}")
                    
                    for key, value in dados_template.items():
                        if key in paragraph.text:
                            # Substituição mais robusta - reconstrói o parágrafo
                            full_text = paragraph.text
                            if key in full_text:
                                new_text = full_text.replace(key, str(value))
                                if new_text != full_text:
                                    # Limpa o parágrafo e adiciona o novo texto
                                    paragraph.clear()
                                    run = paragraph.add_run(new_text)
                                    substituicoes_realizadas[key] = value
                                    logger.info(f"Substituído: '{key}' → '{value}'")
            
            # Itera sobre todas as tabelas do documento (se houver)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            original_text = paragraph.text
                            if '{{' in original_text:
                                logger.info(f"Célula com variáveis encontrada: {original_text}")
                                
                                for key, value in dados_template.items():
                                    if key in paragraph.text:
                                        # Substituição mais robusta - reconstrói o parágrafo
                                        full_text = paragraph.text
                                        if key in full_text:
                                            new_text = full_text.replace(key, str(value))
                                            if new_text != full_text:
                                                # Limpa o parágrafo e adiciona o novo texto
                                                paragraph.clear()
                                                run = paragraph.add_run(new_text)
                                                substituicoes_realizadas[key] = value
                                                logger.info(f"Substituído em tabela: '{key}' → '{value}'")
            
            document.save(output_filename)
            logger.info(f"Contrato gerado: {output_filename}")
            logger.info(f"Total de substituições realizadas: {len(substituicoes_realizadas)}")
            
            # Log das variáveis que não foram substituídas
            variaveis_nao_substituidas = []
            for key in dados_template.keys():
                if key not in substituicoes_realizadas and key.startswith('{{') and key.endswith('}}'):
                    variaveis_nao_substituidas.append(key)
            
            if variaveis_nao_substituidas:
                logger.warning(f"Variáveis não encontradas no template: {variaveis_nao_substituidas}")
            
            return True
            
        except Exception as e:
            logger.error(f"Erro ao preencher template: {str(e)}")
            return False
    
    def send_contract_via_webhook(self, contract_filename, nome_locatario):
        """
        Envia o contrato via webhook
        """
        try:
            # Lê e codifica o arquivo
            with open(contract_filename, "rb") as docx_file:
                encoded_string = base64.b64encode(docx_file.read()).decode('utf-8')
            
            headers = {
                'Content-Type': 'application/json'
            }
            
            payload = {
                "filename": contract_filename,
                "base64": encoded_string,
                "locatario": nome_locatario,
                "caption": f"Contrato Casa da Ana x {nome_locatario}",
                "mimetype": "document/docx"
            }

            logger.info(f"Enviando contrato via webhook para: {nome_locatario}")
            response = requests.post(self.webhook_url, headers=headers, json=payload, timeout=30)
            response.raise_for_status()
            
            logger.info(f"Resposta do webhook - Status: {response.status_code}")
            
            if response.status_code in [200, 201]:
                logger.info("Contrato enviado com sucesso via webhook!")
                # Remove o arquivo após envio bem-sucedido
                try:
                    os.remove(contract_filename)
                    logger.info(f"Arquivo temporário {contract_filename} removido")
                except Exception as e:
                    logger.warning(f"Não foi possível remover arquivo temporário: {e}")
                return True
            else:
                logger.error(f"Falha ao enviar contrato via webhook. Status: {response.status_code}")
                return False
                
        except requests.exceptions.RequestException as e:
            logger.error(f"Erro na requisição para o webhook: {str(e)}")
            return False
        except Exception as e:
            logger.error(f"Erro inesperado ao enviar contrato: {str(e)}")
            return False
    
    def process_contract(self, dados_locatario):
        """
        Processo completo: gera e envia o contrato
        """
        try:
            # Gera o contrato
            contract_filename = self.generate_contract(dados_locatario)
            
            # Envia via webhook
            success = self.send_contract_via_webhook(contract_filename, dados_locatario['nome_do_locatario'])
            
            return {
                'success': success,
                'filename': contract_filename if success else None,
                'message': 'Contrato gerado e enviado com sucesso!' if success else 'Erro ao enviar contrato'
            }
            
        except Exception as e:
            logger.error(f"Erro no processamento do contrato: {str(e)}")
            return {
                'success': False,
                'filename': None,
                'message': f'Erro: {str(e)}'
            } 
