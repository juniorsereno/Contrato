import os
import base64
import requests # Adicionado para chamadas de API
from docx import Document
from docx.shared import Pt

def obter_dados_locatario():
    """Solicita os dados do locatário ao usuário."""
    dados = {}
    print("Por favor, insira os dados do locatário:")
    dados['{{ nome do locatário }}'] = input("Nome completo: ")
    dados['{{ estado civil }}'] = input("Estado Civil: ")
    dados['{{ nacionalidade }}'] = input("Nacionalidade: ")
    dados['{{ profissao }}'] = input("Profissão: ")
    dados['{{ numero do rg }}'] = input("RG: ")
    dados['{{ numero do cpf }}'] = input("CPF: ")
    dados['{{ telefone celular }}'] = input("Telefone Celular: ")
    dados['{{ e-mail }}'] = input("E-mail: ")
    dados['{{ endereco }}'] = input("Endereço completo: ")
    return dados

def preencher_contrato_docx(template_path, output_path_docx, dados_locatario):
    """Preenche o template DOCX com os dados do locatário."""
    try:
        document = Document(template_path)

        # Itera sobre todos os parágrafos do documento
        for paragraph in document.paragraphs:
            for key, value in dados_locatario.items():
                if key in paragraph.text:
                    # Substitui a variável mantendo a formatação o máximo possível
                    # Esta é uma substituição simples. Para formatação complexa, pode ser necessário mais lógica.
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
        
        # Itera sobre todas as tabelas do documento (se houver)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in dados_locatario.items():
                            if key in paragraph.text:
                                for run in paragraph.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, value)
        
        document.save(output_path_docx)
        print(f"Contrato DOCX preenchido salvo em: {output_path_docx}")
        return True
    except Exception as e:
        print(f"Erro ao preencher o contrato DOCX: {e}")
        return False

def enviar_contrato_por_api(caminho_arquivo_docx, nome_locatario, nome_arquivo_gerado):
    """Lê o arquivo DOCX, converte para Base64 e envia via API."""
    try:
        with open(caminho_arquivo_docx, "rb") as docx_file:
            encoded_string = base64.b64encode(docx_file.read()).decode('utf-8')
    except Exception as e:
        print(f"Erro ao ler e codificar o arquivo DOCX para Base64: {e}")
        return False

    api_url = "https://evolution.criativamaisdigital.com.br/message/sendMedia/criativa-suporte"
    # ATENÇÃO: A API Key está hardcoded. Em um ambiente de produção, considere usar variáveis de ambiente.
    api_key = "D2D6BA530A73-4DF0-8AB3-78BD2C514C12"
    
    headers = {
        'Content-Type': 'application/json',
        'apikey': api_key
    }
    
    payload = {
        "number": "556181435045@s.whatsapp.net", # Número fixo conforme solicitado
        "mediatype": "document",
        "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document", # Mimetype correto para docx
        "fileName": nome_arquivo_gerado, # Nome do arquivo que será enviado
        "media": encoded_string,
        "caption": f"Contrato Casa da Ana x {nome_locatario}" # Legenda dinâmica
    }

    try:
        print(f"Enviando contrato para {payload['number']}...")
        response = requests.post(api_url, headers=headers, json=payload, timeout=30) # Adicionado timeout
        response.raise_for_status()  # Lança uma exceção para respostas de erro HTTP (4xx ou 5xx)
        
        print("Resposta da API:")
        print(f"Status Code: {response.status_code}")
        try:
            print(f"Corpo da Resposta: {response.json()}")
        except requests.exceptions.JSONDecodeError:
            print(f"Corpo da Resposta (não JSON): {response.text}")
            
        if response.status_code == 200 or response.status_code == 201:
            print("Contrato enviado com sucesso via API!")
            return True
        else:
            print(f"Falha ao enviar contrato via API. Status: {response.status_code}")
            return False
            
    except requests.exceptions.RequestException as e:
        print(f"Erro na requisição para a API: {e}")
        return False
    except Exception as e:
        print(f"Erro inesperado ao enviar contrato via API: {e}")
        return False

def main():
    template_contrato = "CONTRATO Casa da Ana.docx"
    
    # Verifica se o template existe
    if not os.path.exists(template_contrato):
        print(f"Erro: O arquivo de template '{template_contrato}' não foi encontrado no diretório atual.")
        print(f"Diretório atual: {os.getcwd()}")
        return

    dados_locatario = obter_dados_locatario()
    
    # Sanitiza o nome do locatário para usar no nome do arquivo
    nome_arquivo_base = "".join(c if c.isalnum() or c in (' ', '_') else '_' for c in dados_locatario['{{ nome do locatário }}']).rstrip()
    nome_arquivo_base = nome_arquivo_base.replace(' ', '_')
    
    output_docx_final = f"CONTRATO_{nome_arquivo_base}.docx"

    if preencher_contrato_docx(template_contrato, output_docx_final, dados_locatario):
        print("Contrato DOCX gerado com sucesso.")
        # Enviar o contrato via API
        nome_locatario_para_api = dados_locatario['{{ nome do locatário }}']
        if enviar_contrato_por_api(output_docx_final, nome_locatario_para_api, output_docx_final):
            print("Processo de envio via API concluído.")
        else:
            print("Falha no processo de envio via API. O arquivo DOCX local foi mantido.")
    else:
        print("Falha ao preencher o contrato DOCX.")

if __name__ == "__main__":
    try:
        import docx
        import requests # Verificar também a dependência requests
        print("Dependências 'python-docx' e 'requests' já estão instaladas.")
        main()
    except ImportError:
        print("Uma ou mais dependências (python-docx, requests) não estão instaladas.")
        print("Tentando instalar dependências...")
        import subprocess
        import sys
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "requests"])
            print("Dependências instaladas com sucesso. Por favor, execute o script novamente.")
        except subprocess.CalledProcessError as e:
            print(f"Falha ao instalar dependências: {e}")
            print("Por favor, instale manualmente: pip install python-docx requests")
